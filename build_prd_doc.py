import os, sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

out_dir = r"C:\Users\eppda\Documents\Codex\2026-09-01\new-chat-11\outputs"
os.makedirs(out_dir, exist_ok=True)
docx_path = os.path.join(out_dir, "SubMate_PRD_및_상세_기능명세서_v1.0.docx")

doc = docx.Document()

# Page setup: US Letter, 1 inch margins
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Colors
COLOR_PRIMARY = RGBColor(15, 23, 42)    # Slate 900
COLOR_SECONDARY = RGBColor(51, 65, 85) # Slate 700
COLOR_MUTED = RGBColor(100, 116, 139)  # Slate 500
COLOR_ACCENT = RGBColor(2, 132, 199)   # Sky 600

def set_font(run, font_name="Arial", size_pt=10, color=RGBColor(51, 65, 85), bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=15, color=COLOR_PRIMARY, bold=True)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=12.5, color=COLOR_SECONDARY, bold=True)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=10.5, color=COLOR_SECONDARY, bold=True)
    return p

def add_body(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, font_name="Arial", size_pt=10, color=COLOR_PRIMARY, bold=True)
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=10, color=COLOR_SECONDARY)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, font_name="Arial", size_pt=10, color=COLOR_PRIMARY, bold=True)
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=10, color=COLOR_SECONDARY)
    return p

def add_callout(doc, title, body_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F1F5F9"/>')
    tcPr.append(shd)
    
    tcMar = parse_xml(r'''
        <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="120" w:type="dxa"/>
            <w:left w:w="180" w:type="dxa"/>
            <w:bottom w:w="120" w:type="dxa"/>
            <w:right w:w="180" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(title)
    set_font(r_t, font_name="Arial", size_pt=10.5, color=COLOR_PRIMARY, bold=True)
    
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r_b = p2.add_run(body_text)
    set_font(r_b, font_name="Arial", size_pt=9.5, color=COLOR_SECONDARY)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(4)

def format_table(table, col_widths, headers, rows_data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = Inches(col_widths[i])
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F8FAFC"/>')
        tcPr.append(shd)
        tcMar = parse_xml(r'''
            <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:w="100" w:type="dxa"/>
                <w:left w:w="120" w:type="dxa"/>
                <w:bottom w:w="100" w:type="dxa"/>
                <w:right w:w="120" w:type="dxa"/>
            </w:tcMar>
        ''')
        tcPr.append(tcMar)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_font(run, font_name="Arial", size_pt=9.5, color=COLOR_PRIMARY, bold=True)
            
    for row_idx, row_vals in enumerate(rows_data):
        row = table.add_row()
        for i, val in enumerate(row_vals):
            cell = row.cells[i]
            cell.text = val
            cell.width = Inches(col_widths[i])
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = parse_xml(r'''
                <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:w="80" w:type="dxa"/>
                    <w:left w:w="120" w:type="dxa"/>
                    <w:bottom w:w="80" w:type="dxa"/>
                    <w:right w:w="120" w:type="dxa"/>
                </w:tcMar>
            ''')
            tcPr.append(tcMar)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                set_font(run, font_name="Arial", size_pt=9.0, color=COLOR_SECONDARY)
                
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(r'''
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="94A3B8"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

# Title Block
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
r_title = p_title.add_run("제품 요구사항 정의서 (PRD) & 상세 기능 명세서 (SRS)")
set_font(r_title, font_name="Arial", size_pt=18, color=COLOR_PRIMARY, bold=True)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(12)
r_sub = p_sub.add_run("20대 타겟 올인원 스마트 구독 관리 및 절약 플랫폼 (가칭: SubMate) | Version 1.0 (최종 기획안)")
set_font(r_sub, font_name="Arial", size_pt=10.5, color=COLOR_MUTED)

# Meta info
p_meta = doc.add_paragraph()
p_meta.paragraph_format.space_before = Pt(0)
p_meta.paragraph_format.space_after = Pt(12)
r_meta = p_meta.add_run("작성자: 리드 서비스 기획자 / PO  |  작성일자: 2026. 09. 02  |  Figma 프로젝트: Figma Design System (8 Screens)")
set_font(r_meta, font_name="Arial", size_pt=9.5, color=COLOR_SECONDARY, italic=True)

# Divider
p_div = doc.add_paragraph()
p_div.paragraph_format.space_before = Pt(0)
p_div.paragraph_format.space_after = Pt(12)
p_div_border = parse_xml(r'''
    <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:bottom w:val="single" w:sz="8" w:space="4" w:color="CBD5E1"/>
    </w:pBdr>
''')
p_div._p.get_or_add_pPr().append(p_div_border)

# ----------------- PART 1. PRD -----------------
add_heading_1(doc, "Part 1. 제품 요구사항 정의서 (PRD)")

add_heading_2(doc, "1. Executive Summary (요약)")
add_body(doc, "현대 20대는 OTT, 음원, 쇼핑 멤버십, AI 툴, 클라우드 등 평균 4~7개 이상의 유료 구독 서비스를 동시에 이용하고 있습니다. 이로 인해 '내가 어디에 얼마를 쓰고 있는지 모른 채 깜빡 결제되는 문제'와 '해지 경로 탐색의 번거로움으로 인한 방치'가 반복되고 있습니다.")
add_body(doc, "본 서비스는 OCR 영수증/문자 자동 인식 기반 10초 등록, 결제일 전(D-3, D-1) 사전 알림, 1초 다이렉트 해지 딥링크, 개인화된 대체 프로모션 환승 추천을 통해 사용자가 구독을 온전히 통제하고 실질적인 비용 절감을 체감할 수 있도록 돕는 올인원 구독 관리 플랫폼입니다.")

add_callout(doc, "Core Value Proposition (핵심 가치 제안)", "등록은 가장 쉽게 (AI OCR 10초) -> 관리는 잊지 않게 (D-3, D-1 사전 알림) -> 해지는 가장 빠르게 (1초 딥링크 직행) -> 소비는 더 알뜰하게 (대체 프로모션 환승)")

add_heading_2(doc, "2. Problem Statement & User Pain Points")
add_bullet(doc, "기존 가계부 앱은 서비스명, 금액, 결제일을 일일이 수기 입력해야 해 초기 이탈률이 60% 이상 발생함.", "1. 등록의 번거로움 (High Entry Barrier): ")
add_bullet(doc, "무료 체험 종료일이나 갱신 결제일을 사전에 인지하지 못해 불필요한 결제가 발생한 뒤 카드 승인 문자를 보고 후회함.", "2. 무의식적 결제 및 후회 (Unintended Spending): ")
add_bullet(doc, "서비스마다 계정 관리 및 멤버십 해지 메뉴가 4~5단계 이상 깊숙이 숨겨져 있어 해지를 포기함.", "3. 해지의 높은 탐색 비용 (Dark Patterns): ")
add_bullet(doc, "연간 결제 할인, 통신사/카드사 결합 혜택, 경쟁 서비스의 '첫 달 100원' 등 최적의 환승 프로모션을 유저가 일일이 찾기 어려움.", "4. 절약 기회 부재 (Lack of Optimization): ")

add_heading_2(doc, "3. Product Vision & Goals")
add_body(doc, "모든 구독의 시작, 관리, 해지를 위해 가장 먼저 찾는 대한민국 1등 구독 매니저", "• Vision: ")
add_body(doc, "주간/월간 절약 체감 액션 수 (W/M Saved Actions) - 사전 알림 후 해지 성공 건수 + 환승 프로모션 클릭 및 등록 건수", "• North Star Metric (핵심 지표): ")

# KPI Table
add_heading_3(doc, "핵심 목표 지표 (KPIs)")
kpi_headers = ["지표 구분", "KPI 정의", "목표치 (Launch + 3M)"]
kpi_widths = [1.5, 3.5, 1.5]
kpi_data = [
    ["획득 (Acquisition)", "온보딩 완료율 (회원가입 -> 첫 구독 1개 이상 등록)", "75% 이상"],
    ["활성화 (Activation)", "AI 영수증/문자 인식 등록 성공률", "92% 이상"],
    ["리텐션 (Retention)", "D-3, D-1 푸시 알림 오픈율 (CTR)", "35% 이상"],
    ["비즈니스 (Business)", "프로모션 카드 클릭률 (CTR) 및 전환율 (CVR)", "CTR 12% / CVR 3.5%"]
]
tbl_kpi = doc.add_table(rows=1, cols=3)
format_table(tbl_kpi, kpi_widths, kpi_headers, kpi_data)

add_heading_2(doc, "4. User Persona (타겟 사용자)")
add_body(doc, "김민수 (24세, 대학교 4학년 / 취업준비생) - 실용적 소비를 중시하며 트렌드에 민감한 디지털 네이티브", "• 인적 사항: ")
add_body(doc, "넷플릭스 프리미엄(17,000원), 유튜브 프리미엄(14,900원), 쿠팡 와우(7,890원), 스포티파이(10,900원), ChatGPT Plus(29,000원) -> 매월 총 79,690원 고정 지출 중", "• 이용 현황: ")
add_body(doc, "매달 나가는 총액을 연간 환산액으로 한눈에 파악하고, 무료 체험 만료 전 알림을 받아 바로 해지하며, 안 보는 OTT는 프로모션 특가로 갈아타고 싶음.", "• 핵심 니즈: ")

# ----------------- PART 2. IA -----------------
add_heading_1(doc, "Part 2. 서비스 정보구조도 (IA: Information Architecture)")

add_body(doc, "본 서비스는 20대 사용자의 사용 패턴을 고려하여 최소한의 뎁스(Depth)로 탐색과 해지가 가능하도록 설계되었습니다.")

ia_headers = ["대메뉴 (1 Depth)", "중메뉴 (2 Depth)", "화면/기능 설명", "화면 ID"]
ia_widths = [1.3, 1.4, 2.8, 1.0]
ia_data = [
    ["0. Auth & Onboarding", "0.1 로그인\n0.2 회원가입\n0.3 초기 서비스 선택", "소셜 1초 로그인 / 둘러보기\n아이디 실시간 유효성 검증, 비밀번호, 닉네임\n콜드 스타트 방지 3초 멀티 셀렉터", "00_Auth_Login\n00_Auth_Register\n01_Onboarding"],
    ["1. 홈 대시보드 (Home)", "1.1 지출 요약\n1.2 결제 임박 리스트\n1.3 Empty State\n1.4 절약 큐레이션", "월간 총액 및 연간 환산 토글\nD-Day 순 정렬, 상위 3개 노출\n구독 0개 등록 시 온보딩 유도 안내\n유저 구독 기반 맞춤 환승 프로모션 캐러셀", "01_Home_Dashboard\n01_Home_Empty_State"],
    ["2. 전체 조회 (List)", "2.1 필터/정렬\n2.2 상태 뱃지\n2.3 구독 상세", "결제일순 / 금액순 / 카테고리별 세그먼트\nTODAY, D-1, 결제 실패 경고, 해지 대기 뱃지\n결제수단, 결제주기, D-3/D-1 알림 토글", "02_Subscription_Full_List\n02_Detail_Direct_Cancel"],
    ["3. 구독 추가 (Add)", "3.1 AI 영수증 OCR\n3.2 결제문자 인식\n3.3 수동 검색", "영수증/결제내역 캡처 이미지 업로드 및 자동 파싱\nSMS/앱푸시 복사 붙여넣기 인식\n표준 구독 DB 검색 및 수기 입력", "03_AI_Smart_Add_Modal"],
    ["4. 해지 지원 (Cancel)", "4.1 다이렉트 딥링크\n4.2 경로별 해지 모달\n4.3 절약 피드백", "제공업체 멤버십 해지 페이지 직행 URL\n웹 직행 / 인앱 대행 / 고객센터 전화 3가지 경로\n해지 완료 후 목록 삭제 & 절약 금액 축하 메시지", "02_Detail_Direct_Cancel\n04_Whatsub_Cancel_Modal"],
    ["5. 혜택 & 프로모션", "5.1 혜택 피드\n5.2 제휴 링크 직행", "OTT 환승, 연간 할인, 통신사 결합 특가 큐레이션\n제휴 트래킹 파라미터 연동 및 내 구독 자동 등록", "05_Promotions_Explore"]
]
tbl_ia = doc.add_table(rows=1, cols=4)
format_table(tbl_ia, ia_widths, ia_headers, ia_data)

# ----------------- PART 3. SRS -----------------
add_heading_1(doc, "Part 3. 상세 기능 명세서 (SRS: System Requirement Specification)")

# 1. Auth
add_heading_2(doc, "1. 회원가입 및 인증 (Auth)")

add_heading_3(doc, "[AUTH-01] 회원가입 (Register)")
add_bullet(doc, "신규 사용자의 기본 계정 정보 생성 및 실시간 유효성 검증.", "기능 목적: ")
add_bullet(doc, "00_Auth_Register (로그인 화면에서 '회원가입' 클릭 시 진입)", "화면 위치: ")
add_bullet(doc, "아이디(5~20자 영문소문자/숫자/_/- 허용, 즉시 중복확인 API 호출하여 '사용 가능한 아이디입니다' 녹색 피드백), 비밀번호(8~16자 영문/숫자/특수문자 마스킹 및 보기 토글), 닉네임(3~10자 한/영)", "입력 항목 및 정책: ")
add_bullet(doc, "1) 아이디/이메일 중복 시 필드 하단 에러 표기, 2) 비밀번호 정규식 미충족 시 가이드 노출, 3) 네트워크 실패 시 재시도 토스트 출력.", "예외 처리: ")

add_heading_3(doc, "[AUTH-02] 간편 로그인 및 둘러보기 (Login)")
add_bullet(doc, "20대 유저의 가입 허들을 낮추기 위한 1초 소셜 로그인 지원.", "기능 목적: ")
add_bullet(doc, "00_Auth_Login", "화면 위치: ")
add_bullet(doc, "Apple 로그인(솔리드 블랙 CTA), 카카오 로그인(1.5px 아웃라인 CTA), Google 로그인, 이메일 로그인, 비로그인 유저를 위한 둘러보기(Guest Mode) 텍스트 링크.", "지원 수단: ")

# 2. Onboarding
add_heading_2(doc, "2. 온보딩 (Onboarding)")

add_heading_3(doc, "[ONB-01] 초기 구독 서비스 멀티 셀렉터 (Cold Start Onboarding)")
add_bullet(doc, "가입 직후 빈 화면을 마주하는 부담(Cold Start)을 해소하고 원터치로 구독 현황 셋업.", "기능 목적: ")
add_bullet(doc, "01_Onboarding_Select_Services (Step 1 of 2)", "화면 위치: ")
add_bullet(doc, "영상/OTT(넷플릭스, 유튜브, 티빙, 디즈니+), 음악/쇼핑/AI(쿠팡와우, 스포티파이, 챗GPT) 인기 카드 그리드 노출. 탭 시 선택 토글 및 하단 스티키 바에 'N개 선택됨 | 예상 월 금액' 실시간 합산 출력.", "상세 동작: ")
add_bullet(doc, "'선택 완료하고 결제일 등록하기' 클릭 시 사용자 DB 일괄 저장 후 홈 이동. 상단 '건너뛰기(Skip)' 클릭 시 즉시 홈 이동.", "종료 조건: ")

# 3. Home
add_heading_2(doc, "3. 메인 홈 대시보드 (Home Dashboard)")

add_heading_3(doc, "[HOME-01] 지출 요약 헤더 및 연간 환산 토글")
add_bullet(doc, "등록된 총 구독 개수 및 이번 달 총 결제 예정액 노출. 금액 탭 시 '월간 지출액 <-> 연간 환산 총액' 인터랙티브 전환.", "상세 동작: ")

add_heading_3(doc, "[HOME-02] 결제 임박 구독 카드 리스트 (상위 3개 노출)")
add_bullet(doc, "결제 예정일(D-Day) 기준 오름차순 정렬. 메인 홈에서는 상위 최대 3개만 노출하며, 4개 이상 보유 시 하단에 '더보기(전체보기) >' 버튼 노출.", "표시 정책: ")
add_bullet(doc, "TODAY(솔리드 블랙 펄스 뱃지), D-1(아웃라인 뱃지), D-12(서브틀 그레이 뱃지). 카드 좌측 스와이프 시 [해지 링크 ->], [알림 일시정지] 퀵 액션 노출.", "뱃지 및 인터랙션: ")

add_heading_3(doc, "[HOME-03] 홈 화면 빈 상태 (Empty State - 기능명세서 J8 반영)")
add_bullet(doc, "등록된 구독 서비스가 0개일 때 01_Home_Empty_State 자동 렌더링.", "트리거 조건: ")
add_bullet(doc, "'등록된 구독 서비스가 없습니다' 안내 + [+ 첫 구독 서비스 등록하기] Primary CTA 노출. 하단 프로모션 영역은 '구독 서비스를 등록하면 관련 프로모션을 확인할 수 있어요' 가이드 배너로 대체.", "UI 구성: ")

add_heading_3(doc, "[HOME-04] 맞춤 절약 프로모션 캐러셀 (Flow 4)")
add_bullet(doc, "사용자 보유 구독과 매칭되는 최신 특가/환승 프로모션을 가로 스크롤 미니 카드로 노출 (예: 넷플릭스 유저 대상 '왓챠 첫 달 100원', 디즈니+ 월 결제자 대상 '연간 16% 할인'). 탭 시 제휴 랜딩페이지 직행.", "상세 동작: ")

# 4. Add Subscription
add_heading_2(doc, "4. 구독 추가 및 AI 인식 (Add Subscription)")

add_heading_3(doc, "[ADD-01] AI 영수증/스크린샷 OCR 자동 인식")
add_bullet(doc, "03_AI_Smart_Add_Modal", "화면 위치: ")
add_bullet(doc, "카드 결제 영수증 캡처 사진 업로드 시 AI OCR 엔진이 텍스트 파싱 -> 크롤링된 표준 구독 DB와 매칭하여 서비스명, 요금제, 결제금액, 결제주기 자동 입력. 사용자는 검증 후 [저장 및 알림 설정] 원터치 완료.", "상세 동작: ")
add_bullet(doc, "인식 실패 또는 신뢰도 미달 시 수동 검색창을 팝업하고 인식된 키워드를 자동 입력하여 보정 유도.", "예외 처리: ")

# 5. Detail & Cancel
add_heading_2(doc, "5. 구독 상세 및 다이렉트 해지 (Detail & Cancel)")

add_heading_3(doc, "[DETAIL-01] 구독 상세 조회 및 사전 알림 설정")
add_bullet(doc, "02_Detail_Direct_Cancel", "화면 위치: ")
add_bullet(doc, "결제수단, 결제주기, 다음 결제일 확인. D-3, D-1 사전 푸시 알림 개별 토글 스위치 제공. 하단에 해당 서비스 전용 절약 기회 배너 상시 노출.", "상세 동작: ")

add_heading_3(doc, "[DETAIL-02] 다이렉트 해지 딥링크 및 경로별 가이드")
add_bullet(doc, "'다이렉트 해지 페이지 바로가기 ->' 버튼 탭 시 해당 플랫폼의 '멤버십 해지 설정' 최심단 URL 딥링크로 인앱 브라우저 호출. 해지 완료 후 앱 복귀 시 [해지 완료 확인 및 목록 삭제] 팝업 노출.", "상세 동작: ")

# 6. Full List & Promotions
add_heading_2(doc, "6. 전체 조회 및 프로모션 전용 페이지")

add_heading_3(doc, "[LIST-01] 전체 구독 목록 조회 (02_Subscription_Full_List)")
add_bullet(doc, "결제일 임박순 / 금액 높은순 / 카테고리별 상단 세그먼트 필터 제공. 결제 실패 경고, 해지 대기, D-DAY 뱃지 완비.", "상세 동작: ")

add_heading_3(doc, "[PROMO-01] 맞춤 혜택 & 프로모션 피드 (05_Promotions_Explore)")
add_bullet(doc, "전체 혜택, OTT 환승, 연간 할인, 통신사 결합 카테고리 필터 칩 제공. 제휴 트래킹 파라미터가 포함된 외부 신청 페이지 연결 및 내 구독 자동 등록 체크박스 지원.", "상세 동작: ")

# ----------------- PART 4. BENCHMARK -----------------
add_heading_1(doc, "Part 4. 벤치마크 분석 및 서비스 적용 전략")

bm_headers = ["서비스명", "출신 국가 / 성격", "핵심 시그니처 UI / 기능", "SubMate 서비스 적용 전략"]
bm_widths = [1.2, 1.3, 2.0, 2.0]
bm_data = [
    ["Bobby", "글로벌 1위\n구독 가계부", "극도의 미니멀리즘 카드 UI\n하단 고정 월간/연간 서머리 바", "20대 타겟에 최적화된 군더더기 없는 흑백 미니멀 카드 레이아웃 채택"],
    ["Rocket Money", "미국 1위\n재정 관리 툴", "구독료 인상 감지 알림(Price Hike)\n원클릭 자동 해지 컨시어지 에이전트", "기습 요금 인상 감지 시 대체재 제안 알림 및 1클릭 해지 플로우 도입"],
    ["왓섭 (Whatsub)", "국내 1위\n구독 관리 앱", "월간 결제 캘린더 뷰\n구독 다이어트 (미사용 구독 탐색)\n경로별 해지 안내 모달", "Page 2 레퍼런스로 완벽 구현. 월간 달력 기반 결제일 매핑 및 3가지 해지 경로 제공"],
    ["토스 (Toss)", "국내 핀테크\n고정지출 관리", "카드/계좌 마이데이터 자동 연동\n통신비, 헬스장, 보험료 일괄 분류", "Phase 2~3 로드맵에서 디지털 구독을 넘어 1인 가구 고정지출 전체로 확장"],
    ["피클플러스", "국내 1위\n구독 쉐어링", "계정 공유 1/N 파티원 자동 매칭\n에스크로 기반 자동 N분의 1 정산", "Phase 3 비즈니스 모델로 구독 파티 매칭 및 수수료 모델 연계"]
]
tbl_bm = doc.add_table(rows=1, cols=4)
format_table(tbl_bm, bm_widths, bm_headers, bm_data)

# ----------------- PART 5. ROADMAP -----------------
add_heading_1(doc, "Part 5. 릴리즈 로드맵 & 검증 계획")

rd_headers = ["단계", "목표 기간", "핵심 배포 기능", "품질 검증 기준 (DoD)"]
rd_widths = [1.2, 1.1, 2.4, 1.8]
rd_data = [
    ["Phase 1 (MVP)", "M1 ~ M2", "소셜 로그인, 8대 메인 화면, AI 영수증 OCR 파싱, D-Day 알림, 다이렉트 해지", "OCR 파싱 정확도 90% 이상\n딥링크 해지 성공률 98%"],
    ["Phase 2 (Growth)", "M3 ~ M4", "결제 캘린더 뷰, 가격 인상 감지, 환승 프로모션 제휴 수익화", "MAU 5만 달성\n제휴 카드 클릭률(CTR) 10% 돌파"],
    ["Phase 3 (Scale)", "M5 ~ M6", "구독 1/N 파티원 매칭 & 자동 정산, 고정비 통합 매니저", "자체 결제 정산 에스크로 안정성 99.99% 확보"]
]
tbl_rd = doc.add_table(rows=1, cols=4)
format_table(tbl_rd, rd_widths, rd_headers, rd_data)

doc.save(docx_path)
print(f"DONE: {docx_path}")

