
import os, sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

out_dir = r"C:\Users\eppda\Documents\Codex\2026-09-01\new-chat-11\outputs"
os.makedirs(out_dir, exist_ok=True)
docx_path = os.path.join(out_dir, "SubMate_PRD_v2.0_Comprehensive.docx")

doc = docx.Document()

# Page setup: US Letter, 1 inch margins
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Colors
COLOR_PRIMARY = RGBColor(15, 23, 42)    # Slate 900
COLOR_SECONDARY = RGBColor(51, 65, 85) # Slate 700
COLOR_MUTED = RGBColor(100, 116, 139)  # Slate 500
COLOR_ACCENT = RGBColor(2, 132, 199)   # Sky 600

def set_font(run, font_name="Arial", size_pt=10, color=RGBColor(51, 65, 85), bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=15, color=COLOR_PRIMARY, bold=True)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=12.5, color=COLOR_SECONDARY, bold=True)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=10.5, color=COLOR_SECONDARY, bold=True)
    return p

def add_body(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, font_name="Arial", size_pt=10, color=COLOR_PRIMARY, bold=True)
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=10, color=COLOR_SECONDARY)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, font_name="Arial", size_pt=10, color=COLOR_PRIMARY, bold=True)
    r = p.add_run(text)
    set_font(r, font_name="Arial", size_pt=10, color=COLOR_SECONDARY)
    return p

def add_callout(doc, title, body_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F1F5F9"/>')
    tcPr.append(shd)
    
    tcMar = parse_xml(r'''
        <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="120" w:type="dxa"/>
            <w:left w:w="180" w:type="dxa"/>
            <w:bottom w:w="120" w:type="dxa"/>
            <w:right w:w="180" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(title)
    set_font(r_t, font_name="Arial", size_pt=10.5, color=COLOR_PRIMARY, bold=True)
    
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r_b = p2.add_run(body_text)
    set_font(r_b, font_name="Arial", size_pt=9.5, color=COLOR_SECONDARY)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(4)

def format_table(table, col_widths, headers, rows_data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = Inches(col_widths[i])
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F8FAFC"/>')
        tcPr.append(shd)
        tcMar = parse_xml(r'''
            <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:w="100" w:type="dxa"/>
                <w:left w:w="120" w:type="dxa"/>
                <w:bottom w:w="100" w:type="dxa"/>
                <w:right w:w="120" w:type="dxa"/>
            </w:tcMar>
        ''')
        tcPr.append(tcMar)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_font(run, font_name="Arial", size_pt=9.5, color=COLOR_PRIMARY, bold=True)
            
    for row_idx, row_vals in enumerate(rows_data):
        row = table.add_row()
        for i, val in enumerate(row_vals):
            cell = row.cells[i]
            cell.text = val
            cell.width = Inches(col_widths[i])
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = parse_xml(r'''
                <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:w="80" w:type="dxa"/>
                    <w:left w:w="120" w:type="dxa"/>
                    <w:bottom w:w="80" w:type="dxa"/>
                    <w:right w:w="120" w:type="dxa"/>
                </w:tcMar>
            ''')
            tcPr.append(tcMar)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                set_font(run, font_name="Arial", size_pt=9.0, color=COLOR_SECONDARY)
                
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(r'''
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="94A3B8"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

# Title Block
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
r_title = p_title.add_run("제품 요구사항 정의서 (PRD)")
set_font(r_title, font_name="Arial", size_pt=18, color=COLOR_PRIMARY, bold=True)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(10)
r_sub = p_sub.add_run("20대 타겟 AI 스마트 구독 관리 및 절약 플랫폼 (SubMate) | Version 2.0 (기능명세서 v0.1 완벽 동기화)")
set_font(r_sub, font_name="Arial", size_pt=10.5, color=COLOR_MUTED)

# Meta info
p_meta = doc.add_paragraph()
p_meta.paragraph_format.space_before = Pt(0)
p_meta.paragraph_format.space_after = Pt(12)
r_meta = p_meta.add_run("작성자: 리드 서비스 기획자 / PO  |  문서 상태: 승인 완료 (Approved)  |  기반 문서: RE. (상세 기능명세서) (1).xlsx")
set_font(r_meta, font_name="Arial", size_pt=9.5, color=COLOR_SECONDARY, italic=True)

# Divider
p_div = doc.add_paragraph()
p_div.paragraph_format.space_before = Pt(0)
p_div.paragraph_format.space_after = Pt(12)
p_div_border = parse_xml(r'''
    <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:bottom w:val="single" w:sz="8" w:space="4" w:color="CBD5E1"/>
    </w:pBdr>
''')
p_div._p.get_or_add_pPr().append(p_div_border)

# ----------------- SECTION 1 -----------------
add_heading_1(doc, "1. Executive Summary & Product Overview")
add_body(doc, "SubMate(구독메이트)는 20대 다중 구독 이용자를 위한 AI 기반 구독 관리 및 소비 최적화 플랫폼입니다. 사용자가 이용 중인 다양한 구독 서비스(OTT, 음원, 쇼핑 멤버십, AI 툴, 클라우드 등)를 AI 영수증/SMS 인식으로 간편하게 등록하고, 결제 전 사전 알림(D-3, D-1)을 통해 불필요한 지출을 통제하며, 다이렉트 해지 가이드 및 스마트 대체 프로모션 큐레이션을 통해 실질적인 비용 절감을 실현합니다.")

add_callout(doc, "핵심 가치 제안 (Core Value Proposition)", "1) 초간편 등록: 영수증/문자 AI 인식으로 10초 만에 구독 추가\n2) 지출 사전 통제: D-3, D-1 푸시 알림 및 결제일 갱신 확인 팝업\n3) 다크패턴 해지 해소: 멤버십 해지 페이지 1초 직행 딥링크 & 맞춤 가이드\n4) 스마트 대체 추천: 비슷한 서비스 특가(100원), 무료체험, 연간 할인 3단계 추천")

add_heading_2(doc, "1.1 문제 정의 및 해결 방안 (Problem & Solution)")
ps_headers = ["순번", "사용자 Pain Point", "기존 시장의 한계", "SubMate의 해결책"]
ps_widths = [0.6, 1.8, 2.0, 2.1]
ps_data = [
    ["P1", "등록의 높은 피로도", "서비스명, 금액, 결제일을 수기 입력하여 온보딩 이탈률(60%+) 발생", "AI OCR 및 결제 문자 정규식 파싱으로 10초 자동 등록"],
    ["P2", "무의식적 자동 결제", "무료체험 종료일이나 갱신 결제일을 잊어 결제 후 통보 수신", "D-3, D-1 사전 스마트 알림 및 결제일 경과 갱신 확인 팝업"],
    ["P3", "다크패턴 해지 경로", "서비스마다 해지 메뉴가 4~5단계 숨겨져 있어 해지 포기", "해지 페이지 1초 직행 딥링크 및 동적 해지 가이드 오버레이"],
    ["P4", "비효율적 구독 지출", "연간 할인, 통신사 결합, 100원 특가 정보를 놓쳐 정가 결제", "보유 구독 기반 3단계 스마트 대체 추천 알고리즘 적용"]
]
tbl_ps = doc.add_table(rows=1, cols=4)
format_table(tbl_ps, ps_widths, ps_headers, ps_data)

# ----------------- SECTION 2 -----------------
add_heading_1(doc, "2. Goals & Success Metrics")
add_body(doc, "월간 절약 체감 액션 수 (Monthly Saved Actions) = 사전 알림 후 해지 완료 건수 + 대체 프로모션 환승 완료 건수", "• North Star Metric (북극성 지표): ")

add_heading_2(doc, "2.1 핵심 성과 지표 (KPIs)")
kpi_headers = ["지표 구분", "KPI 정의", "목표치 (Launch + 3M)"]
kpi_widths = [1.5, 3.5, 1.5]
kpi_data = [
    ["획득 (Acquisition)", "온보딩 완료율 (회원가입 -> 첫 구독 등록)", "75% 이상"],
    ["활성화 (Activation)", "AI 영수증/문자 인식 등록 성공률", "92% 이상"],
    ["리텐션 (Retention)", "D-3, D-1 푸시 알림 오픈율 (CTR)", "35% 이상"],
    ["인게이지먼트", "결제일 경과 구독 유지/해제 응답률", "60% 이상"],
    ["비즈니스 (Business)", "프로모션 카드 클릭률 (CTR) / 전환율 (CVR)", "CTR 12% / CVR 3.5%"]
]
tbl_kpi = doc.add_table(rows=1, cols=3)
format_table(tbl_kpi, kpi_widths, kpi_headers, kpi_data)

# ----------------- SECTION 3 -----------------
add_heading_1(doc, "3. 타겟 페르소나 및 사용자 시나리오")
add_body(doc, "김민수 (24세, 대학생 / 취업준비생) - 실용적 소비를 중시하며 트렌드에 민감한 디지털 네이티브", "• 인적 사항: ")
add_body(doc, "넷플릭스 프리미엄(17,000원), 유튜브 프리미엄(14,900원), 쿠팡 와우(7,890원), 스포티파이(10,900원), ChatGPT Plus(29,000원) -> 매월 총 79,690원 고정 지출 중", "• 이용 현황: ")
add_body(doc, "매달 나가는 총액을 연간 환산액으로 한눈에 파악하고, 무료 체험 만료 전 알림을 받아 바로 해지하며, 안 보는 OTT는 프로모션 특가로 갈아타고 싶음.", "• 핵심 니즈: ")

# ----------------- SECTION 4 -----------------
add_heading_1(doc, "4. 서비스 정보구조도 (IA: Information Architecture)")
ia_headers = ["대메뉴 (1 Depth)", "중메뉴 (2 Depth)", "기능 설명 및 화면 ID", "우선순위"]
ia_widths = [1.3, 1.4, 2.8, 1.0]
ia_data = [
    ["0. Auth & Onboarding", "0.1 랜딩페이지\n0.2 회원가입\n0.3 이메일 로그인", "서비스 설명 페이지\n점진적 유효성 검증 폼 (AUTH-01)\n구글/소셜 로그인 연동 (AUTH-02)", "Must\nMust\nMust"],
    ["1. 메인 화면 (Home)", "1.1 구독 정보요약\n1.2 메인 구독리스트\n1.3 프로모션 리스트", "총지출액 및 구독숫자 표시 (SUBS-01)\n결제일 임박순 3개 노출, 정렬 (SUBS-02)\n3단계 우선순위 추천 캐러셀 (SUBS-03)", "Must\nMust\nShould"],
    ["2. 구독 인식 (Capture)", "2.1 영수증/스크린샷\n2.2 SMS 문자 불러오기", "이미지 OCR AI 인식 및 파싱 (RECO-01)\n결제 문자 메시지 텍스트 파싱 (RECO-02)", "Must\nShould"],
    ["3. 구독 추가 (Add)", "3.1 인식 확인/매칭\n3.2 구독서비스 추가\n3.3 단일 항목 수정", "추출 데이터 양식 출력 및 중복확인 (ADDI-01)\nDB 저장 및 성공/실패 토스트 (ADDI-02)\n항목별 개별 수기 수정 (ADDI-03)", "Must\nMust\nShould"],
    ["4. 구독 해지 (Cancel)", "4.1 웹사이트 해지", "해지 페이지 직행 딥링크 및 동적 가이드 (CANC-01)", "Must"],
    ["5. 조회 및 알림", "5.1 D-3, D-1 알림\n5.2 전체 구독 목록\n5.3 구독 갱신 조회", "오전 09:30 사전 결제 푸시 발송 (ALAR-01)\n다차원 정렬/필터링 조회 (ALAR-02)\n결제일 경과 유지/해제 팝업 (ALAR-03)", "Should\nShould\nShould"],
    ["6. 프로모션 추천", "6.1 맞춤형 큐레이션", "카테고리별 특가/환승 피드 (PROM-01)", "Should"]
]
tbl_ia = doc.add_table(rows=1, cols=4)
format_table(tbl_ia, ia_widths, ia_headers, ia_data)

# ----------------- SECTION 5 -----------------
add_heading_1(doc, "5. 상세 기능 요구사항 (Detailed Functional Specifications)")

# 5.1 Auth
add_heading_2(doc, "5.1 로그인 및 인증 (0. 로그인 / 인증)")

add_heading_3(doc, "[AUTH-01] 회원가입 (점진적 검증 폼) - Priority: Must")
add_bullet(doc, "아이디 -> 비밀번호 -> 비밀번호 확인 -> 닉네임 순으로 이전 단계 유효성 조건 만족 시 순차적으로 다음 입력란 표시. 기존 입력란 값은 초기화되지 않고 유지됨. 모든 단계 완료 시 '가입 완료' 버튼 활성화.", "• 점진적 폼 노출: ")
add_bullet(doc, "1) 아이디(5~20자 영문소문자/숫자/_/-): 사용가능 시 '사용 가능한 아이디입니다.' 초록색 표시, 형식 오류 시 빨간색 안내, 중복 시 '중복된 아이디가 있습니다.' 빨간색 표시. 2) 비밀번호(8~16자 영문/숫자/특수문자): 유효 시 '사용가능한 비밀번호 입니다' 초록색 표시. 3) 비밀번호 확인: 불일치 시 '비밀번호 확인을 위해 한번 더 입력해주십시오' 빨간색 표시, 일치 시 '동일한 비밀번호입니다' 초록색 표시. 4) 닉네임: 한/영 3~10자.", "• 입력/출력 규격: ")
add_bullet(doc, "중복 아이디 방지 트랜잭션, 정규식 실시간 검증, 네트워크 장애 시 재시도 토스트 출력.", "• 예외 처리: ")

add_heading_3(doc, "[AUTH-02] 이메일 / 구글 로그인 - Priority: Must")
add_bullet(doc, "구글 로그인 버튼 클릭 시 구글 OAuth 화면 이동 -> 계정 선택 후 로그인 -> 계정 데이터 수신 및 세션 생성.", "• 상세 동작: ")

# 5.2 Main
add_heading_2(doc, "5.2 메인 화면 (01. 메인 화면)")

add_heading_3(doc, "[SUBS-01] 구독서비스 정보조회 - Priority: Must")
add_bullet(doc, "등록된 활성 구독의 총지출액과 구독 숫자를 상단에 표시. 구독 서비스가 없을 때는 '조회 내역이 없습니다.' 문구 표시 및 구독 추가(+) 버튼 강조.", "• 상세 동작 및 예외: ")

add_heading_3(doc, "[SUBS-02] 구독서비스 목록 조회 - Priority: Must")
add_bullet(doc, "결제 예정일 기준 최신(임박)건부터 정렬(알파벳순, 가나다순 정렬 옵션 추가). 메인에는 상위 최대 3개만 노출하며, 3개 이상일 시 하단 '더보기' 버튼 클릭 시 전체 목록으로 이동.", "• 상세 동작: ")
add_bullet(doc, "구독서비스명, 상품명, 결제예정일, 결제액, 서비스 로고 이미지, D-Day 뱃지.", "• 출력 항목: ")
add_bullet(doc, "1) 0개일 때 Empty State ('구독 중인 서비스가 없습니다.' + 구독 추가하기 블록 + 더보기 버튼 숨김), 2) 1~2개일 때 개수만큼 노출, 3) 4개 이상일 때 상위 3개 노출, 4) 이미지 로딩 실패 시 Default Placeholder 노출, 5) 긴 텍스트 말줄임(...) 처리, 6) D-Day(결제일 당일 강조 뱃지), 결제 실패(경고 뱃지), 해지 예정(만료 대기 뱃지), 말일(31일)/윤달(2월 28~29일) 결제일 자동 계산.", "• 예외 처리: ")

add_heading_3(doc, "[SUBS-03] 프로모션 리스트 - Priority: Should")
add_bullet(doc, "사용자 등록 구독 기반 프로모션 API 호출 -> 일치하는 유효 프로모션 선별 -> 최신 등록순 정렬 -> 클릭 시 상세/외부 웹페이지 이동.", "• 상세 동작: ")
add_bullet(doc, "1순위(비슷한 서비스의 더 저렴한 구독제) > 2순위(무료 혹은 이벤트 중인 서비스) > 3순위(연간 결제 등 동일 서비스 할인 요금제).", "• 추천 노출 순서: ")
add_bullet(doc, "등록 구독 부재 시 '구독 서비스를 등록하면 관련 프로모션을 확인할 수 있어요.' 안내, API 지연/실패 시 재시도 버튼, 중복 프로모션 제거, 종료된 프로모션 즉시 제외.", "• 예외 처리: ")

# 5.3 Capture & Add
add_heading_2(doc, "5.3 구독 인식 및 추가 (02 & 03. 구독 인식 및 추가)")

add_heading_3(doc, "[RECO-01 / RECO-02] 영수증 스크린샷 및 SMS 문자 AI 인식 - Priority: Must / Should")
add_bullet(doc, "첨부된 스크린샷 또는 입력된 결제 문자를 AI가 인식하여 양식(서비스명, 플랜, 결제금액, 결제 주기, 다음결제일, 결제 방법)에 맞게 데이터 파싱. 누락 데이터는 크롤링 DB 조회하여 자동 보완 매칭.", "• 상세 동작: ")
add_bullet(doc, "인식 불가 시 '인식오류' 문구 출력 후 메인 화면 복귀.", "• 예외 처리: ")

add_heading_3(doc, "[ADDI-01 / ADDI-02 / ADDI-03] 데이터 확인, 추가 및 수동 수정 - Priority: Must / Should")
add_bullet(doc, "추출 데이터를 폼에 자동 입력하고 기존 DB와 중복 매칭 확인. 중복 확인 시 '중복확인' 문구 출력 후 메인 이동. [구독 추가] 탭 시 DB 저장. 양식 수정 필요 시 항목별 개별 수동 수정 지원(ADDI-03). DB 입력/수정 실패 시 '추가실패'/'수정실패' 출력.", "• 상세 동작 및 예외: ")

# 5.4 Cancel
add_heading_2(doc, "5.4 구독서비스 취소 (04. 구독서비스 취소)")

add_heading_3(doc, "[CANC-01] 구독서비스 취소 및 다이렉트 가이드 - Priority: Must")
add_bullet(doc, "[웹사이트에서 해지하기] 버튼 클릭 -> 해당 서비스 공식 사이트 이동 -> 화면 하단에 단계별 해지 가이드 오버레이 표기 -> 화면 전환에 맞게 가이드 자동 변경 -> 해지 완료 시 해당 서비스 목록에서 즉시 제거.", "• 상세 동작: ")
add_bullet(doc, "사이트 연결 실패 시 '사이트에 연결하지 못했습니다.' 표시. DB에 해지 URL 부재 시 해지 버튼 비활성화(회색 처리).", "• 예외 처리: ")

# 5.5 Alert & List
add_heading_2(doc, "5.5 구독 조회 및 사전결제 알림 (05. 조회 및 알림)")

add_heading_3(doc, "[ALAR-01] D-3, D-1 사전 결제 스마트 푸시 알림 - Priority: Should")
add_bullet(doc, "지정된 시간(기본 오전 09:30 KST)에 푸시 발송 (D-3: '3일 뒤 넷플릭스 17,000원이 결제될 예정입니다.', D-1 무료체험: '내일 티빙 무료체험이 종료되고 13,500원이 결제됩니다.'). 푸시 클릭 시 구독 상세(02_Detail_Direct_Cancel)로 딥링크 이동하며 [웹사이트에서 해지하기] 버튼 하이라이트.", "• 상세 동작: ")
add_bullet(doc, "알림 권한 거부 시 홈 상단 배너 노출, 해지 대기 건 자동 제외, 야간(20:00~08:00) 발송 차단.", "• 예외 처리: ")

add_heading_3(doc, "[ALAR-02] 전체 구독 목록 조회 - Priority: Should")
add_bullet(doc, "홈에서 '전체보기 >' 또는 '구독' 탭 클릭 시 진입. 상단 세그먼트 필터(결제일 임박순 기본, 금액 높은순, 활성/무료체험/해지대기 필터). 카드 탭 시 구독 상세 이동. 당겨서 새로고침 지원.", "• 상세 동작 및 예외: ")

add_heading_3(doc, "[ALAR-03] 구독 목록 갱신 조회 (결제일 경과 확인) - Priority: Should")
add_bullet(doc, "결제일이 지난 서비스가 있을 경우, 결제일 경과 후 처음 앱을 실행할 때 홈 화면 하단에 해당 서비스 유지/해제 여부 확인 팝업 호출. 결정에 따라 다음 결제 주기로 자동 롤오버 또는 목록에서 제거.", "• 상세 동작: ")

# 5.6 Promo
add_heading_2(doc, "5.6 프로모션 탐색과 스마트 대체 추천 (06. 프로모션 추천)")

add_heading_3(doc, "[PROM-01] 맞춤형 프로모션 매칭 및 큐레이션 피드 - Priority: Should")
add_bullet(doc, "사용자 보유 카테고리 기반 유효 프로모션 필터링. 3가지 유형 큐레이션(유사 서비스 환승 특가, 요금제 최적화/연간 전환, 통신사/카드사 결합). 05_Promotions_Explore 화면에서 카테고리 필터 칩 제공 및 외부 제휴 링크 이동.", "• 상세 동작 및 예외: ")

# ----------------- SECTION 6 -----------------
add_heading_1(doc, "6. 예외 처리 매트릭스 (Exception Matrix)")
ex_headers = ["기능 분류", "예외 시나리오", "시스템 대응 및 UI 피드백"]
ex_widths = [1.2, 2.0, 3.3]
ex_data = [
    ["회원가입", "아이디 중복 / 형식 불일치", "'중복된 아이디가 있습니다.' / '아이디 5~20자의 영문 소문자, 숫자와 특수기호만 사용 가능합니다.' 빨간색 에러 표시"],
    ["회원가입", "비밀번호 확인 불일치", "'비밀번호 확인을 위해 한번 더 입력해주십시오' 빨간색 표시 및 닉네임 입력란 숨김 유지"],
    ["메인화면", "등록 구독 0건 (Empty State)", "'구독 중인 서비스가 없습니다.' 안내 + 구독 추가하기 블록 렌더링 + 더보기 버튼 비활성화"],
    ["메인화면", "서비스 이미지 로드 실패", "Default Placeholder / Fallback 아이콘 대체 표시"],
    ["구독인식", "AI OCR 파싱 실패", "'인식오류' 문구 출력 후 메인 화면으로 이동 및 수동 검색 안내"],
    ["구독추가", "기존 구독과 중복 데이터", "'중복확인' 문구 출력 후 메인 화면으로 이동"],
    ["구독추가", "DB 저장 / 수정 실패", "'추가실패' / '수정실패' 문구 출력 및 기존 입력값 보존"],
    ["구독해지", "외부 해지 사이트 연결 실패", "'사이트에 연결하지 못했습니다.' 표시 및 이전 화면 복귀"],
    ["구독해지", "DB에 해지 URL 부재", "[웹사이트에서 해지하기] 버튼 비활성화(Disabled) 및 회색 처리"],
    ["결제알림", "알림 권한 거부 상태", "홈 상단에 '결제 전 알림을 받으려면 알림 권한을 켜주세요' 안내 배너 노출"],
    ["프로모션", "조건 부합 프로모션 0건", "'현재 구독 중인 서비스의 프로모션이 없습니다.' 표시 (미구독 서비스 프로모션 노출 차단)"],
    ["프로모션", "이벤트 조기 종료", "클릭 시 유효기간 재확인 후 이동 차단 및 '종료된 프로모션입니다.' 표시"]
]
tbl_ex = doc.add_table(rows=1, cols=3)
format_table(tbl_ex, ex_widths, ex_headers, ex_data)

# ----------------- SECTION 7 -----------------
add_heading_1(doc, "7. 릴리즈 마일스톤 및 배포 계획")
rd_headers = ["단계", "목표 기간", "포함 기능 명세 (기능 ID)", "완료 정의 (DoD)"]
rd_widths = [1.2, 1.1, 2.4, 1.8]
rd_data = [
    ["Phase 1 (MVP)", "M1 ~ M2", "AUTH-01, AUTH-02, SUBS-01, SUBS-02, RECO-01, ADDI-01, ADDI-02, CANC-01, ALAR-01, ALAR-03", "AI OCR 파싱 성공률 90% 이상\n해지 딥링크 연동 성공률 98%"],
    ["Phase 2 (Growth)", "M3 ~ M4", "SUBS-03, RECO-02, ADDI-03, ALAR-02, PROM-01, 결제 캘린더 뷰", "MAU 5만 달성\n프로모션 CTR 10% 이상 달성"],
    ["Phase 3 (Scale)", "M5 ~ M6", "구독 1/N 파티원 매칭, 고정비 통합 관리 (통신/보험/렌탈)", "자체 정산 시스템 안정성 99.99% 확보"]
]
tbl_rd = doc.add_table(rows=1, cols=4)
format_table(tbl_rd, rd_widths, rd_headers, rd_data)

doc.save(docx_path)
print("SUCCESS:", docx_path)

